
from docx import Document
from docx.shared import Inches
import analisisTemporadas.Modelos.Enumeraciones as en
import pandas as pd
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # Agregar esta línea para importar qn
directorio = 'Agregar dirección para guardarse el documento'
documento = directorio + 'Analisis.docx'

class crearDocumento:
    def __init__(self,archivo = documento):
        self.archivo = archivo
        try:
            # Intentar abrir el documento existente
            self.doc = Document(archivo)
            print(f"Se encontró el archivo '{archivo}' y se cargó.")
        except:
            # Crear un nuevo documento si no existe
            self.doc = Document()
            self.doc.add_heading('Analisis',level=1)
            self.doc.add_paragraph('    En este documento, se analizarán distintas incidencias relevantes de la temporada elegida. Se mostrará la diferencia de los trys recibidos y de los marcados, el analisis de las formaciones fijas, como el scrum y los lineouts, la eficiencia de los pases, los rucks y mauls formados')

            self.doc.save(archivo)
            print(f"No se encontró el archivo '{archivo}', se creó uno nuevo.")
    


    def guardarDocumento(self):
        self.doc.save(documento)


    def existeTitulo(self,Titulo,parrafo):
        lista = []
        if Titulo not in lista:
            lista.append(Titulo)
            doc.crearTitulo(Titulo)
            doc.crearParrafo(parrafo)
     
            
    def crearParrafo(self,texto):
        self.doc.add_paragraph(texto)
        self.doc.save(documento)

    def crearTitulo(self,titulo):
        self.doc.add_heading(titulo,level=1)
        self.doc.save(documento)

    def crearSubtitulo(self,subtitulo):
       
        self.doc.add_heading(subtitulo, level=2) 
        self.doc.save(documento)   
        

    def crearTabla(self,dataframe):
      
        #tablaADocumento(dataframe,doc)
        table = self.doc.add_table(rows=1, cols=len(dataframe.columns))
        table.style = 'Table Grid'
        
        # Agregar encabezados de columna a la tabla
        for col_index, col_name in enumerate(dataframe.columns):
            cell = table.cell(0, col_index)
            cell.text = col_name
            cell.paragraphs[0].alignment = 1  # Alinear al centro
            shading_el = OxmlElement("w:shd")
            shading_el.set(qn("w:fill"), "A9A9A9")  # Color de fondo
            cell._tc.get_or_add_tcPr().append(shading_el)
        
        # Agregar filas de datos a la tabla
        for _, row in dataframe.iterrows():
            row_cells = table.add_row().cells
            for col_index, value in enumerate(row):
                row_cells[col_index].text = str(value)

        self.doc.save(documento)

doc = crearDocumento()

def generarParrafo(Incidencia,maximaIncidenciaPorFecha,maximaIncidenciaPorTemporada,tendenciaTemporada):
  
    if Incidencia == en.TipoIncidencia.Kick.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Clasificacion'])
        
        doc.crearTitulo('Kick')
        parrafo1 = 'En este esta sección, se analizará la evolución de los kicks a lo largo de la temporada. Podrá observarse cuál fue el tipo de kick que pasó más veces, un detalle de cuál fue el que más sucedio en cada fecha, y la tendencia ordenada por la clasificación de los kicks.\n'
        parrafo2 =' A lo largo de la temporada, {0} fue el que más veces terminó sucediendo. A continuación, se mostrará el detalle de cuál fue el tipo de kick que sucedió mayoritariamente \n'.format(maximaIncidenciaPorTemporada)
        
        parrafo3 = '\n Por último, sse detalla la tendencia de cada tipo de kick a lo largo de la temporada, ordenada por fechas y en orden de clasificación \n'
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)
        


     
    elif Incidencia == en.TipoIncidencia.PaseCalidad.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Calidad'])
        tituloPases = 'Pases'
 
        parrafo1 = 'En este esta sección, se analizará la evolución de la calidad de los pases a lo largo de la temporada. Podrá observarse cuál fue el tipo de pase que más se destaca, además un detalle de cuál fue el que más sucedio en cada fecha, y la tendencia ordenada por la clasificación de los pases.\n'
        parrafo2 =' A lo largo de la temporada, {0} fue el que más veces terminó sucediendo. A continuación, se mostrará el detalle de cuál fue la calidad de pase que sucedió mayoritariamente \n'.format(maximaIncidenciaPorTemporada)
        
        parrafo3 = '\n Por últismo, sse detalla la tendencia de cada tipo de calidad de los pases a lo largo de la temporada, ordenada por fechas y en orden de clasificación \n'
        parrafo0 = '    En esta sección, se mostrará el análisis relacionado con los lineouts de la temporada, enfocandose en dos ejes: El resultado y la calidad. El primero de ellos, como su nombre lo indica, se enofcará en cuántos lines se ganaron y cuántos se perdieron, mientras que por el otro lado se mostrará cuántos fueron limpios y cuántos desorganizados'
        doc.existeTitulo(tituloPases,parrafo0)
        doc.crearTitulo('Pase Calidad')
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)
        
    elif Incidencia == en.TipoIncidencia.PaseResultado.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Resultado'])
       
        parrafo1 = 'En este esta sección, se analizará la evolución del resultado de los pases a lo largo de la temporada. Podrá observarse cuál fue el fin de los pases que más se destaca, además un detalle de cuál fue el que más sucedio en cada fecha, y la tendencia ordenada por la clasificación de los pases.\n'
        parrafo2 =' A lo largo de la temporada, {0} fue el que más veces terminó sucediendo. A continuación, se mostrará el detalle de cuál fue el resultado de los pases que sucedió mayoritariamente \n'.format(maximaIncidenciaPorTemporada)
        
        parrafo3 = '\n Por último, sse detalla la tendencia de cada tipo de calidad de los pases a lo largo de la temporada, ordenada por fechas y en orden de clasificación \n'
        
        doc.crearTitulo('Pase Resultado')
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)



    elif Incidencia == en.TipoIncidencia.PuntuacionEnContra.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Percentil'])
        doc.crearTitulo('Puntuación en contra')
        parrafo1 = 'En este esta sección, se analizará la evolución los trys recibidosa lo largo de la temporada. Podrá observarse cuál en qué cuartil del partido se recibió más puntos durante la temporada.\n'
        parrafo2 =' A lo largo de la temporada, {0} el centil que más se recibió puntos. A continuación, se mostrará el detalle de cuál fue el percentil que más puntos sucedió mayoritariamente \n'.format(maximaIncidenciaPorTemporada)
        
        parrafo3 = '\n Por último, sse detalla la tendencia de cada tipo de calidad de los pases a lo largo de la temporada, ordenada por fechas y en orden de clasificación \n'
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)


    elif Incidencia == en.TipoIncidencia.Puntuacion.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Percentil'])
        parrafo0 = 'En esta sección, se analizarán los trys de los encuentros, tanto los hechos por el propio equipo como por los recibidos. Para hacer un mejor análisis, se dividió al tiempo de los encuentros en 4 partes de 20 minutos cada uno, para mostrar en qué parte del partido ocurren más veces esta incidencia'
        parrafo1 = 'En esta primera parte, se analizará la evolución los trys marcados lo largo de la temporada. Podrá observarse cuál en qué cuartil del partido se recibió más puntos durante la temporada.\n'
        parrafo2 =' A lo largo de la temporada, {0} el centil que más trys se marcó. A continuación, se mostrará el detalle de cuál fue el percentil que más puntos sucedió mayoritariamente \n'.format(maximaIncidenciaPorTemporada)
        
        parrafo3 = '\n Por último, sse detalla la tendencia de cada tipo de calidad de los pases a lo largo de la temporada, ordenada por fechas y en orden de clasificación \n'
        doc.existeTitulo('Puntuación',parrafo0)
        doc.crearSubtitulo('Puntuación a favor')

        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)
    
    elif Incidencia == en.TipoIncidencia.LineoutCalidad.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Clasificacion'])
        #    calidadPasesPorcentajeDataframe = pd.DataFrame(pasesPorcentajeCalidad,columns=["Precisos","Imprecisos"],index=index)
        
        parrafo0 = '    En esta sección, se mostrará el análisis relacionado con los lineouts de la temporada, enfocandose en dos ejes: El resultado y la calidad. El primero de ellos, como su nombre lo indica, se enofcará en cuántos lines se ganaron y cuántos se perdieron, mientras que por el otro lado se mostrará cuántos fueron limpios y cuántos desorganizados'
        parrafo1 = '    En esta primera  parte, mostraremos cuál fue el porcentaje de los lineouts ganados y de los perdidos, haciendo énfasis en la evolución y cuál fue el que más sucedio en cada fecha'
        parrafo2 = '    A lo largo de la temporada, los lineouts {0} fueron el tipo de calidad de lineout que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorTemporada)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución del resultado.'
        tituloLineout = 'Line-out'
        
        doc.existeTitulo(tituloLineout,parrafo0)
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)
    
    elif Incidencia == en.TipoIncidencia.LineoutResultado.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Resultado'])
        #    calidadPasesPorcentajeDataframe = pd.DataFrame(pasesPorcentajeCalidad,columns=["Precisos","Imprecisos"],index=index)
        
        parrafo1 = '    En esta segunda parte, mostraremos cuál fue el porcentaje de los lineouts desorganizados y de los limpios, haciendo énfasis en la evolución y cuál fue el que más sucedio en cada fecha'
        parrafo2 = '    A lo largo de la temporada, los lineouts {0} fueron el resultado de lineout que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorTemporada)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución de la calidad.'
        tituloLineout = 'Line-out'
        
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)

    elif Incidencia == en.TipoIncidencia.ScrumResultado.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Resultado'])
            
        parrafo0 = '    En esta sección, se mostrará el análisis relacionado con los scrums de la temporada, enfocandose en dos ejes: El resultado y la calidad. El primero de ellos, como su nombre lo indica, se enofcará en cuántos scrums se ganaron y cuántos se perdieron, mientras que por el otro lado se mostrará cuántos fueron limpios y cuántos desorganizados'
        parrafo1 = '    En esta primera parte, mostraremos cuál fue el porcentaje de los lineouts desorganizados y de los limpios, haciendo énfasis en la evolución y cuál fue el que más sucedio en cada fecha'
        parrafo2 = '    A lo largo de la temporada, los scrum {0} fueron que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorTemporada)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución del resultado.'
        tituloScrum = 'Scrum'
        
        doc.existeTitulo(tituloScrum,parrafo0)
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)


    elif Incidencia == en.TipoIncidencia.ScrumCalidad.value:
        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Calidad'])
            
        parrafo1 = '    En esta segunda parte, mostraremos cuál fue el porcentaje de los scrums desorganizados y de los limpios, haciendo énfasis en la evolución y cuál fue el que más sucedio en cada fecha'
        parrafo2 = '    A lo largo de la temporada, los scrum {0} fueron el tipo de calidad de scrum que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorTemporada)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución de la calidad.'
        tituloScrum = 'Scrum'
        
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)


    elif Incidencia ==  en.TipoIncidencia.TackleResultado.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Resultado'])

        parrafo1 = '    En esta primera parte, mostraremos cuál fue el porcentaje de los tackles ofensivos, los neutrales y los pasivos, haciendo énfasis en la evolución y cuál fue el que más sucedio en cada fecha'
        parrafo2 = '    A lo largo de la temporada, {0} fue el resultado de los tackles  que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorFecha)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución del resultado de los tackles.'
        tituloTackles = 'Tackles'
        
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)

    elif Incidencia == en.TipoIncidencia.TackleEficiencia.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Tipo de Eficiencia'])
            
        parrafo0 = '    En esta sección, se mostrará el análisis relacionado con los tackles de la temporada, enfocandose en dos ejes: El resultado y la eficiencia. El primero de ellos, como su nombre lo indica, se enofcará en cuántos lines se ganaron y cuántos se perdieron, mientras que por el otro lado se mostrará cuántos fueron limpios y cuántos desorganizados'
        parrafo1 = '    En esta segunda parte, mostraremos cuál fue el porcentaje de la eficiencia de los tackles, cuántos fueron realizados y cuántos errados, haciendo énfasis en la evolución y cuál fue el que más sucedió en cada fecha'
        parrafo2 = '    A lo largo de la temporada, {0} fue el tipo de calidad de scrum que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorFecha)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución de la eficiencia.'
        tituloTackles = 'Tackles'
        
        doc.existeTitulo(tituloTackles,parrafo0)
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)

    elif Incidencia == en.TipoIncidencia.RuckAndMaulCalidad.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Calidad'])
            
        parrafo0 = '    En esta sección, se mostrará el análisis relacionado con las formaciones móviles de la temporada, enfocandose en dos ejes: El resultado y la calidad. El primero de ellos, como su nombre lo indica, se enofcará en cuántos rucks y mauls se ganaron y cuántos se perdieron, mientras que por el otro lado se mostrará cuántos fueron rápidos y cuántos lentos'
        parrafo1 = '    En esta primera análisis de los rucks y los mauls, mostraremos cuál fue el porcentaje de calidad de los mismos, cuántos fueron rápidos, lentos, además del porcentaje de penalizados, haciendo énfasis en la evolución y cuál fue el que más sucedió en cada fecha'
        parrafo2 = '    A lo largo de la temporada, {0} fue el tipo de calidad de scrum que más veces sucedió. Además, se dará de forma detalla cuál fue la clasificación que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorFecha)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución de la calidad.'
        tituloTackles = 'Ruck and Maul'
        
        doc.existeTitulo(tituloTackles,parrafo0)
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)


    elif Incidencia == en.TipoIncidencia.RuckAndMaukResultado.value:

        maximaIncidenciaPorFechaDataframe = pd.DataFrame(list(maximaIncidenciaPorFecha.items()), columns=['Fecha', 'Resultado'])
            
        parrafo1 = '    En esta primera análisis de los rucks y los mauls, mostraremos cuál fue el porcentaje de calidad de los mismos, cuántos fueron rápidos, lentos, además del porcentaje de penalizados, haciendo énfasis en la evolución y cuál fue el que más sucedió en cada fecha'
        parrafo2 = '    A lo largo de la temporada, {0} fue el tipo de resultado de las formaciones fijas que más veces sucedió. Además, se dará de forma detalla cuál fue la tipo de resultado que más veces se repitió, pero en cada una de las fechas'.format(maximaIncidenciaPorFecha)
        parrafo3 = '    Por último, se va a detallar, de forma ordenada en cada fecha, la evolución de cada resultado.'
        tituloTackles = 'Tackles'
        
        doc.crearSubtitulo(Incidencia)
        doc.crearParrafo(parrafo1)
        doc.crearParrafo(parrafo2)
        doc.crearTabla(maximaIncidenciaPorFechaDataframe)
        doc.crearParrafo(parrafo3)
        doc.crearTabla(tendenciaTemporada)
        #
        #porcentajeResultadoDataFrame = pd.DataFrame(porcentajeResultado,columns=["Ganados","Perdidos","Penalizados"],index=index)

 
 
        

def generarParrafoEfectividad(porcentajeVictorias):
    parrafo = "Aqui evaluaremos la efectividad de los partidos ganados durante la temporada. En total, de todos los encuentros disponibles se ganó el {0} de los mismos".format(porcentajeVictorias)
    doc.crearTitulo('Efectividad')
    doc.crearParrafo(parrafo)
    

